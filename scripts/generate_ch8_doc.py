"""
Script tạo tài liệu Word CHƯƠNG 8: CÀI ĐẶT VÀ TRIỂN KHAI HỆ THỐNG
Tự động nhúng Hình 8.1 (Swagger OpenAPI) và Hình 8.2 (Kiến trúc Cloud-Native)
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"D:\MyProjects\lms-ai\docs\design\CHUONG_8_CAI_DAT_VA_TRIEN_KHAI.docx"
SWAGGER_IMG = r"C:\Users\lehon\.gemini\antigravity-ide\brain\b10a3cdf-ab31-465a-a135-ff3cc889c80c\swagger_api_ui_1788351148995.jpg"
DEPLOY_IMG = r"C:\Users\lehon\.gemini\antigravity-ide\brain\b10a3cdf-ab31-465a-a135-ff3cc889c80c\deployment_arch_diagram_1788351162673.jpg"

def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def add_figure(doc, img_path, fig_num, caption_text, width_inches=6.2):
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(width_inches))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(14)
        run_cap = p_cap.add_run(f"Hình {fig_num}: {caption_text}")
        run_cap.bold = True
        run_cap.italic = True
        run_cap.font.size = Pt(10.5)
        run_cap.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    else:
        print(f"[WARN] Image missing: {img_path}")

def main():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        
    # Set base styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    
    # ── Chapter Title ──
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(16)
    r_title = p_title.add_run("CHƯƠNG 8: CÀI ĐẶT VÀ TRIỂN KHAI HỆ THỐNG")
    r_title.bold = True
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    # ── 8.1 ──
    p_81 = doc.add_paragraph()
    p_81.paragraph_format.space_before = Pt(12)
    p_81.paragraph_format.space_after = Pt(6)
    r_81 = p_81.add_run("8.1 Cấu hình Môi trường Thực thi và Cấu trúc Mã nguồn")
    r_81.bold = True
    r_81.font.size = Pt(13)
    r_81.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    p = doc.add_paragraph(
        "Hệ thống được xây dựng theo kiến trúc tách biệt (Decoupled Architecture) giữa tầng Giao diện người dùng "
        "(Next.js 14 App Router, TypeScript, TailwindCSS) và tầng Xử lý logic nghiệp vụ phía máy chủ (NestJS Framework, Prisma ORM). "
        "Việc phân tách độc lập này giúp tối ưu hóa hiệu năng tải trang, đảm bảo tính mô-đun hóa cao và thuận tiện cho việc nâng cấp, "
        "bảo trì cũng như mở rộng quy mô độc lập giữa Client và Server."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    
    # ── 8.1.1 ──
    p_811 = doc.add_paragraph()
    p_811.paragraph_format.space_before = Pt(8)
    p_811.paragraph_format.space_after = Pt(4)
    r_811 = p_811.add_run("8.1.1 Cấu hình Biến Môi trường (.env)")
    r_811.bold = True
    r_811.font.size = Pt(12)
    
    p = doc.add_paragraph(
        "Nhằm đảm bảo an toàn thông tin theo nguyên tắc cấu hình tập trung (The Twelve-Factor App), toàn bộ thông số kết nối cơ sở dữ liệu, "
        "khóa bí mật mã hóa và khóa định danh dịch vụ bên thứ ba đều được cô lập trong file môi trường .env phía máy chủ, "
        "hoàn toàn không đẩy lên hệ thống quản lý phiên bản (Git):"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    
    env_items = [
        ("DATABASE_URL: ", "Chuỗi kết nối bảo mật qua giao thức SSL (postgresql://...) tới hệ quản trị CSDL Neon Serverless PostgreSQL trên nền tảng đám mây AWS, tích hợp sẵn kênh kết nối trung gian tự động (Connection Pooling) giúp tối ưu hóa số lượng kết nối đồng thời và chống nghẽn tài nguyên."),
        ("JWT_SECRET & JWT_EXPIRES_IN: ", "Khóa bí mật dùng cho thuật toán chữ ký số HMAC-SHA256 nhằm mã hóa và phát hành mã Token xác thực cho 4 vai trò người dùng trong hệ thống (thời hạn hiệu lực mặc định là 24 giờ)."),
        ("GEMINI_API_KEY: ", "Khóa cấp quyền truy cập Google AI Studio để hệ thống gửi yêu cầu và nhận kết quả phân tích từ mô hình ngôn ngữ lớn Google Gemini phục vụ 3 chức năng GenAI cốt lõi.")
    ]
    for k, v in env_items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.space_after = Pt(4)
        rk = bp.add_run(k)
        rk.bold = True
        bp.add_run(v)
        
    # ── 8.1.2 ──
    p_812 = doc.add_paragraph()
    p_812.paragraph_format.space_before = Pt(10)
    p_812.paragraph_format.space_after = Pt(4)
    r_812 = p_812.add_run("8.1.2 Cấu trúc Tổ chức Mã nguồn Dự án")
    r_812.bold = True
    r_812.font.size = Pt(12)
    
    p = doc.add_paragraph("Mã nguồn dự án được tổ chức phân cấp rõ ràng theo các thư mục chức năng chuyên biệt:")
    p.paragraph_format.space_after = Pt(6)
    
    tree_text = (
        "etc-english-center/\n"
        "├── backend/                  # Tầng xử lý nghiệp vụ Backend (NestJS Framework)\n"
        "│   ├── src/\n"
        "│   │   ├── auth/             # Module xác thực, phân quyền JWT & Roles Guards\n"
        "│   │   ├── users/            # Module quản lý người dùng, hồ sơ 54 học viên & 10 giảng viên\n"
        "│   │   ├── courses/          # Module quản lý danh mục khóa học chuẩn CEFR\n"
        "│   │   ├── classes/          # Module quản lý lớp học, sĩ số tối đa và xếp lịch đào tạo\n"
        "│   │   ├── enrollments/      # Module ghi danh, lập hóa đơn và quản lý phiếu thu học phí\n"
        "│   │   ├── attendances/      # Module điểm danh buổi học (Có mặt / Đi muộn / Vắng mặt)\n"
        "│   │   ├── grades/           # Module bảng điểm quy chế (20% Chuyên cần, 30% Giữa kỳ, 50% Cuối kỳ)\n"
        "│   │   ├── stats/            # Module tổng hợp số liệu báo cáo thống kê & Dashboard\n"
        "│   │   └── ai/               # Module tích hợp Google Gemini GenAI (Prompting & Service)\n"
        "│   └── prisma/               # Schema định nghĩa CSDL 3NF (14 bảng) và file nạp Seed data\n"
        "├── frontend/                 # Tầng giao diện người dùng (Next.js 14 App Router)\n"
        "│   ├── app/\n"
        "│   │   ├── admin/            # Phân hệ giao diện dành cho Quản trị viên\n"
        "│   │   ├── teacher/          # Phân hệ giao diện dành cho Giảng viên\n"
        "│   │   ├── staff/            # Phân hệ giao diện dành cho Nhân viên tư vấn\n"
        "│   │   └── student/          # Phân hệ giao diện dành cho Học viên\n"
        "│   ├── components/           # Thư viện UI tái sử dụng (Design System Dark Theme)\n"
        "│   └── middleware.ts         # Bộ định tuyến bảo vệ và kiểm soát quyền truy cập RBAC\n"
        "└── docs/                     # Toàn bộ tài liệu phân tích, sơ đồ UML/ERD và HDSD"
    )
    
    # Table box for code
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_background(cell, "F1F5F9")
    p_code = cell.paragraphs[0]
    p_code.paragraph_format.space_before = Pt(6)
    p_code.paragraph_format.space_after = Pt(6)
    r_code = p_code.add_run(tree_text)
    r_code.font.name = 'Consolas'
    r_code.font.size = Pt(9.5)
    r_code.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── 8.2 ──
    p_82 = doc.add_paragraph()
    p_82.paragraph_format.space_before = Pt(12)
    p_82.paragraph_format.space_after = Pt(6)
    r_82 = p_82.add_run("8.2 Khởi tạo Cơ sở Dữ liệu và Nạp Dữ liệu Mẫu (Database Seeding)")
    r_82.bold = True
    r_82.font.size = Pt(13)
    r_82.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    p = doc.add_paragraph(
        "Cơ sở dữ liệu của hệ thống được lưu trữ trực tiếp trên nền tảng đám mây Neon Serverless PostgreSQL (đặt tại hạ tầng máy chủ AWS), "
        "đảm bảo tính sẵn sàng cao (High Availability), tự động sao lưu dữ liệu và co giãn tài nguyên theo nhu cầu thực tế."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph("Quy trình tạo lập và nạp dữ liệu mẫu được tự động hóa hoàn toàn thông qua công cụ Prisma ORM:")
    p.paragraph_format.space_after = Pt(6)
    
    # Steps
    p_s1 = doc.add_paragraph()
    p_s1.paragraph_format.space_after = Pt(4)
    r = p_s1.add_run("• Bước 1: Đồng bộ cấu trúc bảng (Schema Synchronization):")
    r.bold = True
    
    p = doc.add_paragraph(
        "Thực thi lệnh npx prisma db push để tạo lập tự động 14 bảng quan hệ chuẩn hóa 3NF trên Neon Cloud, "
        "đồng thời thiết lập đầy đủ khóa chính (Primary Key), khóa ngoại (Foreign Key), chỉ mục (Index) và các ràng buộc toàn vẹn dữ liệu (Unique, Check Constraint, Cascading)."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    
    p_s2 = doc.add_paragraph()
    p_s2.paragraph_format.space_after = Pt(4)
    r = p_s2.add_run("• Bước 2: Nạp dữ liệu kiểm thử (Database Seeding):")
    r.bold = True
    
    p = doc.add_paragraph(
        "Thực thi lệnh npx prisma db seed để nạp sẵn danh mục dữ liệu mẫu toàn diện phục vụ cho công tác kiểm thử và nghiệm thu, bao gồm:"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    
    sub_seed = [
        ("Tài khoản mẫu cho 4 vai trò: ", "admin01 (Quản trị), 10 Giảng viên (teacher01 → teacher10), 2 Tư vấn viên (staff01, staff02), và 54 Học viên (student01 → student54). Toàn bộ mật khẩu mặc định đều được mã hóa bằng thuật toán băm an toàn Argon2 trước khi lưu vào CSDL."),
        ("Danh mục đào tạo: ", "6 Khóa học chuẩn CEFR (A1, A2, B1, B2, C1, IELTS Master), danh sách phòng học và 6 lớp học phân bổ theo các ca học trong tuần."),
        ("Kịch bản sĩ số tối đa & Phân chia phiếu thu: ", "Lớp IELTS-B1-01 đạt sĩ số tối đa 25/25 học viên (100% tải), lịch sử phiếu thu được phân tách độc lập giữa các tư vấn viên (staff01 phụ trách 33 phiếu, staff02 phụ trách 21 phiếu)."),
        ("Dữ liệu đánh giá học tập: ", "Toàn bộ bảng điểm kết thúc khóa của 54 học viên được tính toán tự động theo đúng tỷ lệ quy chế đào tạo, cho ra kết quả nghiệm thu rõ ràng: 40 học viên ĐẠT (87.0%), 6 học viên KHÔNG ĐẠT (13.0%) và 5 học viên ĐANG HỌC.")
    ]
    for k, v in sub_seed:
        bp = doc.add_paragraph(style='List Bullet')
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.left_indent = Inches(0.5)
        bp.paragraph_format.space_after = Pt(4)
        rk = bp.add_run(k)
        rk.bold = True
        bp.add_run(v)

    # ── 8.3 ──
    p_83 = doc.add_paragraph()
    p_83.paragraph_format.space_before = Pt(12)
    p_83.paragraph_format.space_after = Pt(6)
    r_83 = p_83.add_run("8.3 Đóng gói, Tài liệu hóa API và Triển khai Hệ thống (Deployment)")
    r_83.bold = True
    r_83.font.size = Pt(13)
    r_83.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    p = doc.add_paragraph(
        "Hệ thống được thiết kế và triển khai hoàn chỉnh theo mô hình kiến trúc Cloud-Native, tận dụng tối đa các nền tảng dịch vụ đám mây chuyên dụng để tách biệt độc lập giữa tầng Giao diện, tầng Dịch vụ và tầng Dữ liệu."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    
    # ── 8.3.1 ──
    p_831 = doc.add_paragraph()
    p_831.paragraph_format.space_before = Pt(8)
    p_831.paragraph_format.space_after = Pt(4)
    r_831 = p_831.add_run("8.3.1 Tài liệu hóa và Kiểm thử API tự động với Swagger / OpenAPI")
    r_831.bold = True
    r_831.font.size = Pt(12)
    
    p = doc.add_paragraph(
        "Phía Backend (NestJS) được cấu hình tích hợp thư viện @nestjs/swagger, tự động trích xuất toàn bộ đặc tả kỹ thuật của 32 RESTful APIs thành tài liệu tương tác trực tiếp chuẩn OpenAPI 3.0:"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    
    sw_features = [
        ("Tương tác kiểm thử trực quan: ", "Cho phép lập trình viên và hội đồng kiểm thử tra cứu chi tiết các endpoint, cấu trúc tham số (Request Body/Params/Query) và mã phản hồi (HTTP Status Code 200, 201, 400, 401, 403, 409, 500) trực tiếp trên giao diện trình duyệt mà không cần sử dụng công cụ kiểm thử bên ngoài (như Postman)."),
        ("Cơ chế xác thực Bearer Token: ", "Tích hợp nút Authorize cho phép nạp mã định danh JWT Token sau khi đăng nhập, hỗ trợ kiểm thử an toàn các API yêu cầu quyền hạn của từng nhóm tác nhân (Admin, Teacher, Staff, Student).")
    ]
    for k, v in sw_features:
        bp = doc.add_paragraph(style='List Bullet')
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.space_after = Pt(4)
        rk = bp.add_run(k)
        rk.bold = True
        bp.add_run(v)
        
    # Add Figure 8.1
    add_figure(doc, SWAGGER_IMG, "8.1", "Giao diện tài liệu tương tác Swagger OpenAPI 32 RESTful APIs của hệ thống", width_inches=6.2)
    
    # ── 8.3.2 ──
    p_832 = doc.add_paragraph()
    p_832.paragraph_format.space_before = Pt(10)
    p_832.paragraph_format.space_after = Pt(4)
    r_832 = p_832.add_run("8.3.2 Kiến trúc Triển khai Cloud-Native")
    r_832.bold = True
    r_832.font.size = Pt(12)
    
    p = doc.add_paragraph("Hệ thống hoàn chỉnh được đưa vào vận hành trên môi trường mạng thông qua mô hình phân tầng điện toán đám mây:")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    
    cloud_tiers = [
        ("Tầng Frontend (Vercel Edge Network): ", "Ứng dụng Next.js được build và triển khai tự động qua cơ chế CI/CD liên kết với kho mã nguồn GitHub. Mã nguồn giao diện được phân phối qua mạng lưới máy chủ biên (Edge Network) của Vercel, giúp người dùng truy cập hệ thống với độ trễ thấp và tối ưu tốc độ tải trang."),
        ("Tầng Backend (Render / Railway Container): ", "Mã nguồn xử lý nghiệp vụ NestJS được đóng gói dưới dạng Docker Container độc lập, cấu hình tự động lắng nghe và xử lý các yêu cầu RESTful API từ Client với chứng chỉ mã hóa đường truyền bảo mật HTTPS/SSL."),
        ("Tầng Cơ sở dữ liệu (Neon Serverless Cloud): ", "Toàn bộ dữ liệu vận hành được lưu trữ trên cụm đám mây Neon PostgreSQL, hỗ trợ tách rời giữa tầng lưu trữ (Storage) và tầng tính toán (Compute), tự động mở rộng tài nguyên khi có tải cao và duy trì kết nối ổn định 24/7.")
    ]
    for k, v in cloud_tiers:
        bp = doc.add_paragraph(style='List Bullet')
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.space_after = Pt(4)
        rk = bp.add_run(k)
        rk.bold = True
        bp.add_run(v)
        
    # Add Figure 8.2
    add_figure(doc, DEPLOY_IMG, "8.2", "Mô hình kiến trúc triển khai phân tán Cloud-Native của hệ thống", width_inches=6.2)
    
    # ── 8.3.3 ──
    p_833 = doc.add_paragraph()
    p_833.paragraph_format.space_before = Pt(10)
    p_833.paragraph_format.space_after = Pt(4)
    r_833 = p_833.add_run("8.3.3 Thông số Môi trường và Đường dẫn Truy cập Vận hành")
    r_833.bold = True
    r_833.font.size = Pt(12)
    
    p = doc.add_paragraph("Các đường dẫn vận hành chính thức phục vụ công tác kiểm thử và nghiệm thu sản phẩm:")
    p.paragraph_format.space_after = Pt(6)
    
    urls = [
        ("Địa chỉ ứng dụng Giao diện (Frontend Production): ", "https://[tên-miền-dự-án-của-bạn].vercel.app (hoặc http://localhost:3000 trên môi trường phát triển)"),
        ("Địa chỉ máy chủ Dịch vụ (Backend API Base URL): ", "https://[tên-miền-backend].app/api/v1 (hoặc http://localhost:8000/api/v1 trên môi trường phát triển)"),
        ("Địa chỉ Tài liệu tương tác API (Swagger UI): ", "https://[tên-miền-backend].app/api/docs (hoặc http://localhost:8000/api/docs trên môi trường phát triển)")
    ]
    for k, v in urls:
        bp = doc.add_paragraph(style='List Bullet')
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.space_after = Pt(4)
        rk = bp.add_run(k)
        rk.bold = True
        bp.add_run(v)
        
    # Save document
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"SUCCESS: Generated {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
