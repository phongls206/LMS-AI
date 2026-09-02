"""
Script tạo file Word chứa các mục MỚI (7.8, 7.9, 7.10) để gộp vào cuối Chương 7.
Cập nhật CẤU TRÚC MÃ NGUỒN CHUẨN XÁC 100% theo đúng thư mục thực tế của dự án.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"D:\MyProjects\lms-ai\docs\design\CHUONG_7_PHAN_THEM_7.8_7.9_7.10.docx"

def set_cell_bg(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    # ══════════════════════════════════════════════
    # NOTE
    # ══════════════════════════════════════════════
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_note.paragraph_format.space_after = Pt(16)
    r = p_note.add_run("( Phần bổ sung gộp vào cuối Chương 7 — Copy toàn bộ nội dung bên dưới\nvà paste ngay SAU mục 7.7 Ma trận Truy vết trong file báo cáo chính )")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # ══════════════════════════════════════════════
    # 7.8
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("7.8 Cấu hình Môi trường Thực thi và Cấu trúc Mã nguồn")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    p = doc.add_paragraph(
        "Sau khi hoàn tất thiết kế kiến trúc và đặc tả API, hệ thống được cài đặt và cấu hình "
        "môi trường thực thi để sẵn sàng cho giai đoạn lập trình và kiểm thử."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)

    # ── 7.8.1 ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("7.8.1 Cấu hình Biến Môi trường (.env)")
    r.bold = True
    r.font.size = Pt(13)

    p = doc.add_paragraph(
        "Nhằm đảm bảo an toàn thông tin theo nguyên tắc cấu hình tập trung (The Twelve-Factor App), "
        "toàn bộ thông số nhạy cảm đều được cô lập trong file .env phía máy chủ, "
        "hoàn toàn không đẩy lên hệ thống quản lý phiên bản (Git):"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)

    env_items = [
        ("DATABASE_URL", "Chuỗi kết nối bảo mật qua giao thức SSL tới hệ quản trị CSDL Neon Serverless PostgreSQL trên nền tảng đám mây AWS, tích hợp sẵn Connection Pooling giúp tối ưu hóa số lượng kết nối đồng thời."),
        ("JWT_SECRET & JWT_EXPIRES_IN", "Khóa bí mật phục vụ cơ chế xác thực JWT đã trình bày tại Mục 7.3 (thời hạn hiệu lực mặc định 24 giờ)."),
        ("GEMINI_API_KEY", "Khóa cấp quyền truy cập Google AI Studio phục vụ 3 chức năng GenAI đã đặc tả tại Mục 7.2.")
    ]
    for k, v in env_items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.space_after = Pt(3)
        rk = bp.add_run(f"{k}: ")
        rk.bold = True
        bp.add_run(v)

    # ── 7.8.2 ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("7.8.2 Cấu trúc Tổ chức Mã nguồn Dự án")
    r.bold = True
    r.font.size = Pt(13)

    p = doc.add_paragraph(
        "Mã nguồn dự án được tổ chức phân cấp rõ ràng theo các thư mục chức năng chuyên biệt, "
        "phản ánh trực tiếp kiến trúc phân tầng đã thiết kế tại Mục 7.1:"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)

    # CẤU TRÚC CHUẨN XÁC 100% THEO FILE SYSTEM DỰ ÁN
    tree = (
        "lms-ai/                               # Thư mục gốc dự án ETC English Center\n"
        "├── backend/                          # TẦNG BACKEND (NestJS Framework)\n"
        "│   ├── prisma/                       # Cấu hình CSDL & Dữ liệu mẫu\n"
        "│   │   ├── schema.prisma             # Định nghĩa Schema 14 bảng quan hệ chuẩn 3NF\n"
        "│   │   └── seed.ts                   # Kịch bản nạp Seed data (54 HV, 10 GV, Lớp full 25/25)\n"
        "│   └── src/                          # Mã nguồn TypeScript phía Backend\n"
        "│       ├── common/                   # Thành phần dùng chung (Guards, Filters, Decorators)\n"
        "│       ├── prisma/                   # Module kết nối cơ sở dữ liệu PrismaService\n"
        "│       ├── modules/                  # 9 Modules nghiệp vụ chính:\n"
        "│       │   ├── auth/                 # Xác thực đăng nhập, đổi mật khẩu & cấp JWT\n"
        "│       │   ├── users/                # Quản lý tài khoản, hồ sơ học viên & giảng viên\n"
        "│       │   ├── courses/              # Quản lý danh mục khóa học chuẩn CEFR\n"
        "│       │   ├── classes/              # Quản lý lớp học, phòng học, ca học & sĩ số\n"
        "│       │   ├── enrollments/          # Quản lý ghi danh, hóa đơn & phiếu thu học phí\n"
        "│       │   ├── attendances/          # Quản lý điểm danh buổi học (Có mặt/Đi muộn/Vắng)\n"
        "│       │   ├── grades/               # Quản lý bảng điểm quy chế (20% - 30% - 50%)\n"
        "│       │   ├── statistics/           # Tổng hợp báo cáo KPI, doanh thu & tỷ lệ Đạt\n"
        "│       │   └── ai/                   # Tích hợp Google Gemini GenAI (Prompting & Fallback)\n"
        "│       ├── app.module.ts             # Module gốc điều phối toàn bộ ứng dụng\n"
        "│       └── main.ts                   # Khởi tạo Server, cấu hình Swagger & ValidationPipe\n"
        "├── frontend/                         # TẦNG FRONTEND (Next.js 14 App Router)\n"
        "│   └── src/\n"
        "│       ├── app/                      # Bộ định tuyến phân trang theo 4 vai trò RBAC:\n"
        "│       │   ├── admin/                # Phân hệ Quản trị viên (Dashboard, GV, HV, Lớp, Báo cáo)\n"
        "│       │   ├── teacher/              # Phân hệ Giảng viên (Điểm danh, Nhập điểm, Lịch dạy)\n"
        "│       │   ├── staff/                # Phân hệ Tư vấn viên (Thu học phí, Quản lý học viên)\n"
        "│       │   ├── student/              # Phân hệ Học viên (Lịch học, Bảng điểm, Trợ lý AI)\n"
        "│       │   ├── login/                # Màn hình đăng nhập chuyển đổi 4 vai trò\n"
        "│       │   ├── change-password/      # Màn hình đổi mật khẩu người dùng\n"
        "│       │   └── layout.tsx            # Khung giao diện dùng chung (Dark Theme)\n"
        "│       ├── components/               # Thư viện UI tái sử dụng (AppLayout, Sidebar, Modal)\n"
        "│       ├── services/                 # Tầng giao tiếp RESTful API (api.ts - Axios/Fetch Client)\n"
        "│       └── types/                    # Định nghĩa cấu trúc dữ liệu TypeScript DTO\n"
        "├── docs/                             # Toàn bộ tài liệu SRS, sơ đồ UML/ERD và báo cáo\n"
        "├── scripts/                          # Scripts tạo dữ liệu seed và xuất tài liệu kỹ thuật\n"
        "├── .env                              # Tệp biến môi trường bảo mật\n"
        "└── docker-compose.yml                # Tệp cấu hình đóng gói môi trường Container"
    )
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "F1F5F9")
    pc = cell.paragraphs[0]
    pc.paragraph_format.space_before = Pt(6)
    pc.paragraph_format.space_after = Pt(6)
    rc = pc.add_run(tree)
    rc.font.name = 'Consolas'
    rc.font.size = Pt(8.5)
    rc.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════════
    # 7.9
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("7.9 Khởi tạo Cơ sở Dữ liệu và Nạp Dữ liệu Mẫu (Database Seeding)")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    p = doc.add_paragraph(
        "Cơ sở dữ liệu được lưu trữ trên nền tảng đám mây Neon Serverless PostgreSQL (AWS), "
        "đảm bảo tính sẵn sàng cao (High Availability) và tự động sao lưu. "
        "Quy trình tạo lập và nạp dữ liệu mẫu được tự động hóa hoàn toàn thông qua Prisma ORM:"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)

    # Step 1
    bp = doc.add_paragraph(style='List Bullet')
    bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bp.paragraph_format.space_after = Pt(4)
    rk = bp.add_run("Bước 1 — Đồng bộ cấu trúc bảng (Schema Synchronization): ")
    rk.bold = True
    bp.add_run(
        "Thực thi lệnh npx prisma db push để tạo lập tự động 14 bảng quan hệ chuẩn hóa 3NF "
        "(đã thiết kế tại Chương 6), đồng thời thiết lập đầy đủ khóa chính, khóa ngoại, chỉ mục "
        "và các ràng buộc toàn vẹn dữ liệu."
    )

    # Step 2
    bp = doc.add_paragraph(style='List Bullet')
    bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bp.paragraph_format.space_after = Pt(4)
    rk = bp.add_run("Bước 2 — Nạp dữ liệu kiểm thử (Database Seeding): ")
    rk.bold = True
    bp.add_run(
        "Thực thi lệnh npx prisma db seed để nạp sẵn danh mục dữ liệu mẫu toàn diện "
        "phục vụ cho công tác kiểm thử và nghiệm thu, bao gồm:"
    )

    seed_items = [
        ("Tài khoản mẫu cho 4 vai trò: ",
         "admin01 (Quản trị), 10 Giảng viên (teacher01 → teacher10), "
         "2 Tư vấn viên (staff01, staff02), và 54 Học viên (student01 → student54). "
         "Toàn bộ mật khẩu mặc định đều được mã hóa bằng thuật toán Argon2 trước khi lưu vào CSDL."),
        ("Danh mục đào tạo: ",
         "6 Khóa học chuẩn CEFR (A1, A2, B1, B2, C1, IELTS Master), "
         "danh sách phòng học và 6 lớp học phân bổ theo các ca học trong tuần."),
        ("Kịch bản sĩ số tối đa: ",
         "Lớp IELTS-B1-01 đạt sĩ số tối đa 25/25 học viên (100% tải). "
         "Lịch sử phiếu thu được phân tách độc lập giữa các tư vấn viên "
         "(staff01 phụ trách 33 phiếu, staff02 phụ trách 21 phiếu)."),
        ("Dữ liệu đánh giá học tập: ",
         "Toàn bộ bảng điểm kết thúc khóa của 54 học viên được tính toán tự động theo "
         "đúng tỷ lệ quy chế đào tạo (20% Chuyên cần, 30% Giữa kỳ, 50% Cuối kỳ), "
         "cho ra kết quả nghiệm thu: 40 học viên ĐẠT (87.0%), "
         "6 học viên KHÔNG ĐẠT (13.0%) và 5 học viên ĐANG HỌC.")
    ]
    for k, v in seed_items:
        sp = doc.add_paragraph(style='List Bullet 2')
        sp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        sp.paragraph_format.space_after = Pt(3)
        rk = sp.add_run(k)
        rk.bold = True
        sp.add_run(v)

    # ══════════════════════════════════════════════
    # 7.10
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("7.10 Tài liệu hóa API tự động (Swagger) và Đường dẫn Truy cập Vận hành")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    # Swagger
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("7.10.1 Tài liệu hóa và Kiểm thử API với Swagger / OpenAPI")
    r.bold = True
    r.font.size = Pt(13)

    p = doc.add_paragraph(
        "Phía Backend (NestJS) được cấu hình tích hợp thư viện @nestjs/swagger, "
        "tự động trích xuất toàn bộ đặc tả kỹ thuật của 32 RESTful APIs "
        "(đã liệt kê tại Mục 7.6) thành tài liệu tương tác trực tiếp chuẩn OpenAPI 3.0:"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)

    sw = [
        ("Tương tác kiểm thử trực quan: ",
         "Cho phép lập trình viên và hội đồng kiểm thử tra cứu chi tiết các endpoint, "
         "cấu trúc tham số (Request Body/Params/Query) và mã phản hồi "
         "(HTTP Status Code 200, 201, 400, 401, 403, 409, 500) trực tiếp trên giao diện "
         "trình duyệt mà không cần sử dụng công cụ bên ngoài (như Postman)."),
        ("Cơ chế xác thực Bearer Token: ",
         "Tích hợp nút Authorize cho phép nạp mã JWT Token sau khi đăng nhập, "
         "hỗ trợ kiểm thử an toàn các API yêu cầu quyền hạn của từng nhóm tác nhân "
         "(Admin, Teacher, Staff, Student).")
    ]
    for k, v in sw:
        bp = doc.add_paragraph(style='List Bullet')
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.space_after = Pt(3)
        rk = bp.add_run(k)
        rk.bold = True
        bp.add_run(v)

    # URLs
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("7.10.2 Thông số Môi trường và Đường dẫn Truy cập Vận hành")
    r.bold = True
    r.font.size = Pt(13)

    p = doc.add_paragraph(
        "Các đường dẫn vận hành chính thức phục vụ công tác kiểm thử và nghiệm thu sản phẩm:"
    )
    p.paragraph_format.space_after = Pt(6)

    urls = [
        ("Địa chỉ ứng dụng Giao diện (Frontend): ",
         "https://[tên-miền].vercel.app (hoặc http://localhost:3000 trên môi trường phát triển)"),
        ("Địa chỉ máy chủ API (Backend): ",
         "https://[tên-miền-backend].app/api/v1 (hoặc http://localhost:8000/api/v1)"),
        ("Địa chỉ Tài liệu API (Swagger UI): ",
         "https://[tên-miền-backend].app/api/docs (hoặc http://localhost:8000/api/docs)")
    ]
    for k, v in urls:
        bp = doc.add_paragraph(style='List Bullet')
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.space_after = Pt(3)
        rk = bp.add_run(k)
        rk.bold = True
        bp.add_run(v)

    # Save
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"SUCCESS: Updated {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
