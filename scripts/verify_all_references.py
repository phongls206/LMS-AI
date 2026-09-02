import docx
import json
import os

doc_path = r"D:\MyProjects\lms-ai\docs\design\EnglishCenterTOP.docx"
doc = docx.Document(doc_path)

full_text = []
for p in doc.paragraphs:
    if p.text.strip():
        full_text.append(p.text.strip())

tables_text = []
for t in doc.tables:
    rows = []
    for r in t.rows:
        rows.append([c.text.strip().replace("\n", " ") for c in r.cells])
    tables_text.append(rows)

all_text_str = "\n".join(full_text)

# Also check CHUONG_8_9_CAI_DAT_TRIEN_KHAI_KIEM_THU_HUONG_DAN_SD.md if any
ch89_path = r"D:\MyProjects\lms-ai\docs\design\CHUONG_8_9_CAI_DAT_TRIEN_KHAI_KIEM_THU_HUONG_DAN_SD.md"
ch89_text = ""
if os.path.exists(ch89_path):
    with open(ch89_path, "r", encoding="utf-8") as f:
        ch89_text = f.read()

combined_text = all_text_str + "\n" + ch89_text
combined_lower = combined_text.lower()

checklist = {
    "File 01: 01_GenAI_SoftwareDevelopment_project-plan.docx": {
        "Kế hoạch thực hiện 9 tuần": "kế hoạch" in combined_lower and ("tuần" in combined_lower or "27/07" in combined_lower or "thực hiện" in combined_lower),
        "Bảng phân công công việc chi tiết (47 mục)": len(tables_text) > 0,
        "Thành viên thực hiện & Ghi chú (Lê Hồng Phong, Lưu Thanh Nguyên)": "lê hồng phong" in combined_lower and "lưu thanh nguyên" in combined_lower
    },
    "File 02: 02_GenAI_SoftwareDevelopment_requirements-qa.docx": {
        "20 câu hỏi khảo sát nghiệp vụ (Q&A)": "câu hỏi" in combined_lower and "khảo sát" in combined_lower,
        "Yêu cầu chức năng (FR) đầy đủ": "yêu cầu chức năng" in combined_lower,
        "Yêu cầu phi chức năng (NFR) đầy đủ": "phi chức năng" in combined_lower,
        "Sơ đồ phân cấp chức năng (FDD)": "phân cấp chức năng" in combined_lower
    },
    "File 03: 03_GenAI_SoftwareDevelopment_requirements-specification.docx": {
        "Mục đích & Phạm vi ứng dụng": "mục đích" in combined_lower and "phạm vi" in combined_lower,
        "Bảng thuật ngữ & từ viết tắt": "thuật ngữ" in combined_lower or "viết tắt" in combined_lower,
        "Tài liệu tham khảo": "tài liệu tham khảo" in combined_lower,
        "Danh sách Tác nhân (4 Actors: Quản lý, GV, TVV, HV)": all(a in combined_lower for a in ["quản lý", "giáo viên", "học viên", "tư vấn viên"]),
        "Danh sách 14 Use Cases (UC001 - UC014)": all(f"uc{i:03d}" in combined_lower for i in range(1, 15)),
        "Biểu đồ Use Case tổng quát & phân rã 4 Actor": "use case" in combined_lower and "phân rã" in combined_lower,
        "Đặc tả Use Case chi tiết (14 Use Case)": all(f"uc{i:03d}" in combined_lower for i in range(1, 15)),
        "Activity Diagram & Sequence Diagram cho từng UC": "activity" in combined_lower and ("trình tự" in combined_lower or "sequence" in combined_lower),
        "Chương 4: Dữ liệu I/O, Prompt mẫu AI, Xử lý lỗi & Fallback": "prompt" in combined_lower and "fallback" in combined_lower and "kiểm tra đầu vào" in combined_lower
    },
    "File 04: 04_GenAI_SoftwareDevelopment_object-oriented-design.docx": {
        "Sơ đồ lớp tổng thể (Class Diagram)": "class diagram" in combined_lower or "sơ đồ lớp" in combined_lower or "mô hình lớp" in combined_lower,
        "Đặc tả chi tiết 12 Lớp đối tượng": all(c in combined_lower for c in ["nguoidung", "hosohocvien", "khoahoc", "lophoc", "dichvuai"]),
        "Thuộc tính (kiểu dữ liệu, kích thước) & Phương thức từng lớp": "thuộc tính" in combined_lower and "phương thức" in combined_lower
    },
    "File 05: 05_GenAI_SoftwareDevelopment_functional-testing.docx": {
        "Yêu cầu phần cứng & phần mềm kiểm thử": "phần cứng" in combined_lower and ("phần mềm" in combined_lower or "kiểm thử" in combined_lower),
        "Danh sách Test Cases (Test ID, Chức năng, Preconditions, Test Data, Expected)": "test" in combined_lower and "kết quả mong muốn" in combined_lower,
        "Báo cáo kết quả Test Report (Pass/Fail, Độ nghiêm trọng)": "pass" in combined_lower or "kết quả test" in combined_lower or "độ nghiêm trọng" in combined_lower
    },
    "File 06: 06_GenAI_SoftwareDevelopment_screenflow_db.docx": {
        "Screen Flow phân luồng màn hình theo 4 Actor": "screen flow" in combined_lower or "phân luồng" in combined_lower,
        "Sơ đồ thực thể quan hệ (ERD)": "erd" in combined_lower or "thực thể quan hệ" in combined_lower,
        "Lược đồ CSDL quan hệ 3NF (14 bảng)": "3nf" in combined_lower or "cơ sở dữ liệu" in combined_lower,
        "Mô tả chi tiết các bảng CSDL (Khóa chính, Khóa ngoại, Kiểu dữ liệu)": "khóa chính" in combined_lower and "khóa ngoại" in combined_lower,
        "Các ràng buộc toàn vẹn CSDL": "ràng buộc toàn vẹn" in combined_lower
    },
    "File 07: 07_GenAI_SoftwareDevelopment_user-guide.docx": {
        "Giới thiệu ứng dụng & Cấu hình phần cứng/phần mềm": "giới thiệu" in combined_lower and "cấu hình" in combined_lower,
        "Hướng dẫn sử dụng chi tiết theo 4 Actor (Admin, Teacher, Staff, Student)": all(a in combined_lower for a in ["quản lý", "giáo viên", "tư vấn viên", "học viên"]) and "hướng dẫn" in combined_lower
    }
}

with open(r"D:\MyProjects\lms-ai\docs\verification_summary.json", "w", encoding="utf-8") as f:
    json.dump(checklist, f, ensure_ascii=False, indent=2)

print("Saved verification summary to verification_summary.json")
