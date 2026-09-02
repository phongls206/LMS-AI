import docx
import os
import json

ref_dir = r"D:\MyProjects\lms-ai\docs\references"
files = [
    "01_GenAI_SoftwareDevelopment_project-plan.docx",
    "02_GenAI_SoftwareDevelopment_requirements-qa.docx",
    "03_GenAI_SoftwareDevelopment_requirements-specification.docx",
    "04_GenAI_SoftwareDevelopment_object-oriented-design.docx",
    "05_GenAI_SoftwareDevelopment_functional-testing.docx",
    "06_GenAI_SoftwareDevelopment_screenflow_db.docx",
    "07_GenAI_SoftwareDevelopment_user-guide.docx"
]

report = {}

for f in files:
    path = os.path.join(ref_dir, f)
    doc = docx.Document(path)
    headings = []
    paragraphs = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if p.style.name.startswith("Heading") or t.startswith(("#", "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "I.", "II.", "III.", "IV.", "V.", "Chương", "CHƯƠNG")):
            headings.append(t)
        paragraphs.append(t)
    
    tables_summary = []
    for i, t in enumerate(doc.tables):
        first_row = [c.text.strip().replace("\n", " ") for c in t.rows[0].cells] if t.rows else []
        tables_summary.append(f"Table {i+1} ({len(t.rows)} rows x {len(t.columns)} cols): {first_row}")
        
    report[f] = {
        "total_paragraphs": len(paragraphs),
        "total_tables": len(doc.tables),
        "headings": headings[:25],
        "sample_text": paragraphs[:10],
        "tables_summary": tables_summary[:5]
    }

with open(r"D:\MyProjects\lms-ai\docs\references_analysis.json", "w", encoding="utf-8") as out:
    json.dump(report, out, ensure_ascii=False, indent=2)

print("Saved analysis to references_analysis.json")
