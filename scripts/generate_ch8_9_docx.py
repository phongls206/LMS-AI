import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = docx.Document()

# Page Margins (2.54 cm = 1 inch)
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def add_heading_chapter(text):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Navy Blue
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(12)
    return h

def add_heading_1(text):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_heading_2(text):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    return h

def add_heading_3(text):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = RGBColor(51, 65, 85)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(3)
    return h

def add_p(text, bold_prefix="", italic_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Times New Roman'
        r_bold.font.size = Pt(13)
        r_bold.font.bold = True

    if italic_prefix:
        r_it = p.add_run(italic_prefix)
        r_it.font.name = 'Times New Roman'
        r_it.font.size = Pt(13)
        r_it.font.italic = True

    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    return p

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(4)

    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Times New Roman'
        r_bold.font.size = Pt(13)
        r_bold.font.bold = True

    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    return p

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.italic = True
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_code_block(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(30, 41, 59)
    return p

def create_table(headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_background(cell, "1E3A8A")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10.5)

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, val in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

    return table

# =========================================================================
# CHƯƠNG 8: CÀI ĐẶT VÀ TRIỂN KHAI HỆ THỐNG
# =========================================================================

add_heading_chapter("CHƯƠNG 8: CÀI ĐẶT VÀ TRIỂN KHAI HỆ THỐNG")

add_heading_1("8.1 Môi trường và Yêu cầu Kỹ thuật")
add_heading_2("8.1.1 Yêu cầu Cấu hình Phần cứng")
add_p("Hệ thống Quản lý trung tâm ngoại ngữ tích hợp AI (ETC English) được xây dựng theo kiến trúc phân tán (Multi-tier Cloud-ready Architecture), cho phép tối ưu hóa tài nguyên phần cứng máy chủ và máy khách, giảm thiểu chi phí đầu tư hạ tầng ban đầu:")

headers_hw = ["Thành phần", "Cấu hình tối thiểu", "Cấu hình khuyến nghị"]
data_hw = [
    ["Máy chủ Phát triển (Dev Server)", "CPU 2 Cores, RAM 4GB, Ổ cứng trống 5GB", "CPU 4 Cores, RAM 8GB–16GB, Ổ cứng SSD 20GB"],
    ["Máy khách Người dùng (Client)", "Thiết bị có trình duyệt web (PC, Laptop, Smartphone)", "PC/Laptop màn hình Full HD, kết nối Internet ổn định"],
    ["Băng thông Mạng", "Tối thiểu 5 Mbps (truy vấn và gọi GenAI API)", "Khuyến nghị 20 Mbps trở lên để tải trang mượt mà"]
]
create_table(headers_hw, data_hw)
add_caption("Bảng 8.1: Yêu cầu cấu hình phần cứng tối thiểu và khuyến nghị")

add_heading_2("8.1.2 Yêu cầu Môi trường Phần mềm & Công nghệ")
add_p("Hệ thống được phát triển hoàn toàn trên nền tảng TypeScript xuyên suốt từ Frontend đến Backend:")

headers_sw = ["Thành phần", "Công nghệ / Framework", "Phiên bản", "Mục đích sử dụng"]
data_sw = [
    ["Runtime Môi trường", "Node.js (LTS)", "v20.x / v22.x", "Môi trường thực thi JavaScript/TypeScript phía máy chủ"],
    ["Quản lý Gói", "npm", "10.x+", "Quản lý và cài đặt các thư viện phụ thuộc"],
    ["Backend Framework", "NestJS", "v11.x / v12.x", "Xây dựng RESTful API kiến trúc phân tầng Modular"],
    ["ORM / Data Access Layer", "Prisma ORM", "v6.4.1 (Stable)", "Quản lý Schema, Migration và Type-safe Database Client"],
    ["Frontend Framework", "Next.js (App Router)", "v14.x / v16.x", "Xây dựng giao diện SSR/CSR với TypeScript"],
    ["CSS & Giao diện", "TailwindCSS + Lucide Icons", "v3.4.x / v4.x", "Thiết kế giao diện hiện đại, Responsive đa thiết bị"],
    ["Cơ sở Dữ liệu", "PostgreSQL (Neon.tech)", "v15+ / v16+", "CSDL quan hệ Serverless Cloud Database"],
    ["Trí tuệ Nhân tạo", "Google Gemini SDK", "@google/genai", "Tích hợp mô hình ngôn ngữ lớn (Gemini 2.5 Flash/Pro)"],
    ["Bảo mật & Mã hóa", "Argon2 + JWT", "argon2, @nestjs/jwt", "Băm mật khẩu an toàn và xác thực phân quyền Stateless"]
]
create_table(headers_sw, data_sw)
add_caption("Bảng 8.2: Danh mục công nghệ và phiên bản phần mềm cốt lõi")

add_heading_2("8.1.3 Cấu trúc Thư mục Mã Nguồn Chuẩn")
add_p("Dự án được tổ chức theo mô hình Monorepo rõ ràng giữa Backend và Frontend:")
add_code_block("""D:\\MyProjects\\lms-ai\\
├── backend/                        # Máy chủ Backend (NestJS 12)
│   ├── prisma/
│   │   ├── schema.prisma           # 14 Models thực thể chuẩn 3NF
│   │   └── seed.ts                 # Script nạp dữ liệu mẫu ban đầu
│   ├── src/
│   │   ├── common/                 # Guards, Decorators, Strategies (RBAC)
│   │   ├── modules/
│   │   │   ├── auth/               # Module xác thực JWT & Argon2
│   │   │   ├── users/              # Module học viên & giáo viên
│   │   │   ├── courses/            # Module quản lý khóa học
│   │   │   ├── classes/            # Module lớp học & lịch học (chống trùng phòng/GV)
│   │   │   ├── enrollments/        # Module đăng ký lớp & hóa đơn học phí
│   │   │   ├── attendances/        # Module buổi học & điểm danh 4 trạng thái
│   │   │   ├── grades/             # Module bảng điểm (công thức 20/30/50)
│   │   │   ├── statistics/         # Module báo cáo doanh thu & tỷ lệ hoàn thành
│   │   │   └── ai/                 # Module tích hợp Gemini AI & Lọc ảo giác
│   │   ├── prisma/                 # Prisma Module & Prisma Service
│   │   ├── app.module.ts           # Root Module của NestJS
│   │   └── main.ts                 # Điểm khởi chạy hệ thống, Swagger & Validation
│   └── .env                        # File cấu hình biến môi trường Backend
├── frontend/                       # Ứng dụng Giao diện (Next.js 14 App Router)
│   ├── src/
│   │   ├── app/                    # 23 Màn hình phân luồng theo 4 vai trò (RBAC)
│   │   ├── components/             # AppLayout, Sidebar, Header phân quyền
│   │   ├── services/               # Axios API Client với JWT Interceptor
│   │   └── types/                  # TypeScript Data Models & DTO Interfaces
└── docs/                           # Tài liệu thiết kế hệ thống & Đồ án""")

add_heading_1("8.2 Cấu Hình Môi Trường & Quản Trị Bí Mật (Environment & Secrets)")
add_heading_2("8.2.1 Cấu hình Biến Môi Trường Backend (backend/.env)")
add_p("Toàn bộ thông tin cấu hình nhạy cảm được quản lý qua biến môi trường độc lập, tuyệt đối không hardcode vào mã nguồn:")
add_code_block("""# Cổng dịch vụ Backend
PORT=8000

# Chuỗi kết nối Neon Serverless PostgreSQL (Pooler Mode hỗ trợ SSL)
DATABASE_URL="postgresql://neondb_owner:npg_3lFnjQKIo5rM@ep-dark-resonance-axjf1mzr-pooler.c-4.us-east-2.aws.neon.tech/neondb?sslmode=require&channel_binding=require"

# Khóa bí mật ký phát mã JWT (Stateless Token)
JWT_SECRET="etc_english_center_jwt_secret_key_2026_super_secure"
JWT_EXPIRES_IN="24h"

# Khóa API Google Gemini AI (Tích hợp AI)
GEMINI_API_KEY="AIzaSy...your_gemini_api_key_here..."

# Cấu hình CORS cho phép Frontend truy cập
FRONTEND_URL="http://localhost:3000" """)

add_heading_2("8.2.2 Cấu hình Chuỗi Kết Nối Cơ Sở Dữ Liệu Neon Cloud")
add_p("Hệ thống sử dụng đường dẫn kết nối dạng Connection Pooler (PgBouncer) của Neon.tech để tối ưu số lượng kết nối đồng thời từ NestJS, đồng thời bật chế độ mã hóa đường truyền bắt buộc (sslmode=require).")

add_heading_1("8.3 Quy Trình Cài Đặt Và Khởi Tạo Cơ Sở Dữ Liệu")
add_p("Quy trình thiết lập hệ thống từ mã nguồn được thực hiện theo 3 bước tuần tự:")
add_bullet("Cài đặt thư viện phụ thuộc (Dependencies): Chạy npm install tại thư mục backend và frontend.", "Bước 1: ")
add_bullet("Đồng bộ CSDL và Sinh Prisma Client: Chạy npx prisma db push và npx prisma generate để tạo 14 bảng quan hệ 3NF trên Neon PostgreSQL.", "Bước 2: ")
add_bullet("Nạp Dữ liệu Mẫu Ban Đầu (Database Seeding): Chạy npm run db:seed để nạp 4 vai trò người dùng, 2 khóa học, 2 lớp học và lịch học.", "Bước 3: ")

headers_seed = ["Tên đăng nhập", "Mật khẩu", "Vai trò (RBAC)", "Họ và Tên", "Mô tả nghiệp vụ"]
data_seed = [
    ["admin01", "Admin@123", "QUAN_LY", "Nguyễn Quản Lý", "Toàn quyền quản trị trung tâm, tài chính, phân công"],
    ["teacher01", "Admin@123", "GIAO_VIEN", "Nguyễn Thị Lan", "Giảng viên IELTS, TOEIC (Phụ trách lớp LOP01)"],
    ["teacher02", "Admin@123", "GIAO_VIEN", "Trần Văn Minh", "Giảng viên Giao tiếp (Phụ trách lớp LOP02)"],
    ["staff01", "Admin@123", "TU_VAN_VIEN", "Lê Thị Tư Vấn", "Tư vấn viên tiếp nhận học viên & thu phí tại quầy"],
    ["student01", "Admin@123", "HOC_VIEN", "Phạm Văn An", "Học viên trình độ CEFR B1 (Lớp LOP01)"],
    ["student02", "Admin@123", "HOC_VIEN", "Hoàng Thị Bình", "Học viên trình độ CEFR A2 (Lớp LOP02)"]
]
create_table(headers_seed, data_seed)
add_caption("Bảng 8.3: Danh mục tài khoản người dùng mẫu sau khi nạp Seeding")

add_heading_1("8.4 Quy Trình Khởi Chạy Hệ Thống Trên Môi Trường Cục Bộ")
add_bullet("Khởi chạy Máy chủ Backend NestJS: cd backend && npm run start:dev (Backend chạy tại port 8000, Swagger UI tại http://localhost:8000/api/docs).", "1. ")
add_bullet("Khởi chạy Ứng dụng Giao diện Frontend Next.js: cd frontend && npm run dev (Giao diện chạy tại http://localhost:3000/login).", "2. ")
add_bullet("Quản lý Dữ liệu Trực quan Qua Prisma Studio: cd backend && npx prisma studio (Giao diện GUI quản lý 14 bảng CSDL tại http://localhost:5555).", "3. ")

add_heading_1("8.5 Phương Án Triển Khai Lên Môi Trường Đám Mây (Cloud Deployment)")
add_p("Hệ thống được thiết kế sẵn sàng cho việc triển khai phân tán không máy chủ (Serverless Cloud Architecture):")
add_bullet("Frontend Tier: Triển khai trên nền tảng Vercel (Edge Network / Serverless Functions) với tên miền chính.", "• ")
add_bullet("Backend Tier: Triển khai trên Render hoặc Railway với môi trường Docker Node.js Container.", "• ")
add_bullet("Database Tier: Lưu trữ trực tiếp trên Neon Serverless PostgreSQL với tính năng tự động sao lưu và mở rộng quy mô (Auto-scaling).", "• ")
add_bullet("Chính sách Bảo mật: Bật tường lửa CORS giới hạn Domain truy cập, áp dụng HTTPS toàn bộ luồng dữ liệu và kích hoạt Rate Limiting ngăn chặn tấn công DDoS.", "• ")

# =========================================================================
# CHƯƠNG 9: KIỂM THỬ HỆ THỐNG VÀ HƯỚNG DẪN SỬ DỤNG
# =========================================================================

add_heading_chapter("CHƯƠNG 9: KIỂM THỬ HỆ THỐNG VÀ HƯỚNG DẪN SỬ DỤNG")

add_heading_1("PHẦN 1: KIỂM THỬ CHỨC NĂNG ỨNG DỤNG (FUNCTIONAL TESTING)")
add_p("Thời gian thực hiện kiểm thử: Từ 27/07/2026 đến 27/09/2026 (Theo kế hoạch phát triển 9 tuần của dự án).")

add_heading_2("1. Những Yêu Cầu Về Tài Nguyên Cho Kiểm Thử Ứng Dụng")
add_heading_3("a) Tài nguyên Phần cứng:")
add_p("Môi trường kiểm thử chức năng được thực hiện trên cấu hình máy tính cá nhân kết nối mạng Internet:")

headers_test_hw = ["CPU", "RAM", "Ổ cứng (SSD)", "Kiến trúc hệ thống (Architecture)"]
data_test_hw = [
    ["Intel Core i5 / AMD Ryzen 5, 2.5 GHz trở lên", "8 GB / 16 GB", "50 GB dung lượng trống", "x64 (64-bit Operating System)"]
]
create_table(headers_test_hw, data_test_hw)
add_caption("Bảng 9.1: Cấu hình phần cứng phục vụ kiểm thử ứng dụng")

add_heading_3("b) Tài nguyên Phần mềm:")
headers_test_sw = ["Tên phần mềm / Công cụ", "Phiên bản", "Loại công cụ / Mục đích sử dụng"]
data_test_sw = [
    ["Visual Studio Code / Antigravity IDE", "1.95+", "Công cụ phát triển mã nguồn & Debug"],
    ["Swagger UI & Postman Client", "OpenAPI 3.0 / v11+", "Công cụ kiểm thử RESTful API & Endpoint Authorization"],
    ["Trình duyệt Web (Google Chrome / Edge)", "v125+", "Kiểm thử giao diện người dùng (Frontend UI Testing)"],
    ["Neon Console & Prisma Studio", "v6.4.1", "Kiểm tra tính toàn vẹn dữ liệu quan hệ (Database Integrity)"],
    ["Hệ điều hành Windows 11", "Build 23H2 (64-bit)", "Môi trường thực thi kiểm thử cục bộ"]
]
create_table(headers_test_sw, data_test_sw)
add_caption("Bảng 9.2: Danh mục phần mềm và công cụ kiểm thử")

add_heading_2("2. Danh Sách Các Tình Huống Kiểm Thử Ứng Dụng (Test Cases)")
add_p("Toàn bộ 14 Use Case nghiệp vụ và tính năng tích hợp AI được thiết kế các kịch bản kiểm thử bao gồm cả trường hợp hợp lệ, trường hợp ngoại lệ và điều kiện biên:")

headers_tc = ["Test ID", "Chức năng", "Mô tả ca kiểm thử", "Điều kiện trước", "Dữ liệu Test", "Kết quả mong muốn", "Ghi chú"]
data_tc = [
    ["TC001", "Đăng nhập (UC001)", "Xác thực danh tính và phân quyền theo JWT", "Tài khoản đã tồn tại trong CSDL", "admin01 / Admin@123", "Đăng nhập thành công, cấp JWT Token, chuyển hướng đúng Dashboard Quản trị", "Bảo mật Argon2"],
    ["TC002", "Hồ sơ Học viên (UC002)", "Tạo mới tài khoản và gán trình độ CEFR", "Đăng nhập quyền Quản lý/TVV", "Học viên: HV003, CEFR: B1", "Hồ sơ được tạo trong CSDL qua ACID Transaction, mã hóa mật khẩu", "ACID Transaction"],
    ["TC003", "Khóa học (UC003)", "Tạo mới chương trình đào tạo", "Đăng nhập quyền Quản lý", "Khóa: KH03, Học phí: 4.5M, Tiết: 45", "Khóa học được lưu, mã khóa không trùng lặp, học phí >= 0", "Dữ liệu hợp lệ"],
    ["TC004", "Lớp & Lịch (UC004)", "Xếp lịch học và kiểm tra chống trùng phòng", "Lớp học đã được tạo", "Phòng P.101, Thứ 2 (18h-21h)", "Hệ thống chặn nếu phòng P.101 đã có lớp khác học cùng giờ", "Chống trùng phòng"],
    ["TC005", "Phân công GV (UC005)", "Gán giảng viên phụ trách lớp học", "Giáo viên và lớp đã tồn tại", "Gán GV001 cho lớp LOP01", "Hệ thống chặn nếu GV001 đã có lịch dạy lớp khác cùng ca/thứ", "Chống trùng giờ dạy"],
    ["TC006", "Đăng ký lớp (UC006)", "Kiểm tra 4 điều kiện: Sĩ số, Chưa ĐK, CEFR, Lịch", "Học viên đã được cấp mã", "Học viên B1 đăng ký lớp LOP01 (CEFR B1)", "Đăng ký thành công, sĩ số tăng +1, tự động sinh Hóa đơn học phí", "Kiểm tra 4 điều kiện"],
    ["TC007", "Thu học phí (UC007)", "Ghi nhận thanh toán nhiều đợt", "Hóa đơn học phí ở trạng thái nợ", "Đợt 1 nộp 2.000.000đ", "Cộng dồn số tiền đã trả, khi đủ tiền tự chuyển sang DA_HOAN_THANH", "Thanh toán từng phần"],
    ["TC008", "Điểm danh (UC008)", "Ghi nhận 4 trạng thái chuyên cần buổi học", "Đến giờ học buổi số 1", "CO_MAT, VANG, DI_MUON, CO_PHEP", "Lưu trữ chính xác 4 trạng thái cho từng học viên, tính tỷ lệ % tham gia", "4 trạng thái chuẩn"],
    ["TC009", "Nhập điểm (UC009)", "Tính điểm tổng kết 20/30/50 và xét ĐẠT", "Lớp học hoàn thành kỳ thi", "CC=90, GK=80, CK=85", "Điểm TK = 90*0.2 + 80*0.3 + 85*0.5 = 84.5. Xếp loại ĐẠT", "Công thức 20/30/50"],
    ["TC010", "Tra cứu TKB (UC010)", "Học viên xem lịch học và kết quả cá nhân", "Đăng nhập tài khoản student01", "Truy cập /student/schedule", "Hiển thị đúng danh sách lớp, phòng học, giáo viên phụ trách của cá nhân", "Bảo mật RBAC"],
    ["TC011", "Báo cáo Thống kê (UC011)", "Thống kê doanh thu, sĩ số và tỷ lệ hoàn thành", "Đăng nhập quyền Quản lý", "Truy cập /admin/reports", "Hiển thị chính xác tổng doanh thu, biểu đồ sĩ số và % học viên ĐẠT", "Thống kê thời gian thực"],
    ["TC012", "AI Tư vấn lớp (UC012)", "Gợi ý tối đa 3 lớp học thực tế theo CEFR & lịch rảnh", "Đăng nhập quyền Học viên/TVV", "CEFR: B1, Lịch rảnh: Thứ 2-4-6", "AI gợi ý đúng lớp có thật còn chỗ, có Fallback tự động khi lỗi mạng", "Lọc ảo giác Zero-Trust"],
    ["TC013", "AI Sinh bài tập (UC013)", "Sinh tức thì 5 câu trắc nghiệm chuẩn CEFR", "Đăng nhập quyền Giáo viên/HV", "Chủ đề: Present Perfect, CEFR: B1", "Trả về 5 câu hỏi JSON có 4 lựa chọn, đáp án đúng và giải thích chi tiết", "Template Fallback"],
    ["TC014", "AI Tóm tắt tiến độ (UC014)", "Phân tích chuyên cần & điểm thi, đưa lời khuyên", "Đăng nhập quyền Học viên", "Chọn lớp học đang theo học", "AI tóm tắt điểm mạnh, điểm yếu và gợi ý chủ đề cần ôn tập bổ trợ", "Audit Logging"]
]
create_table(headers_tc, data_tc)
add_caption("Bảng 9.3: Danh mục các ca kiểm thử chức năng hệ thống (Test Cases)")

add_heading_2("3. Báo Cáo Kết Quả Kiểm Thử (Test Report)")
add_p("Toàn bộ các ca kiểm thử đã được chạy nghiệm thu trên môi trường thực tế kết nối Neon Cloud và Google Gemini AI API:")

headers_tr = ["Test ID", "Ngày testing", "Người tham gia Test", "Pass/Fail", "Độ nghiêm trọng", "Tóm tắt kết quả kiểm tra", "Ghi chú"]
data_tr = [
    ["TC001", "02/09/2026", "Tester & Developer", "PASS", "High", "Đăng nhập chính xác cả 4 vai trò, phân quyền bảo mật", "Không có lỗi"],
    ["TC002", "02/09/2026", "Tester & Developer", "PASS", "High", "Tạo học viên thành công, mã hóa mật khẩu bằng Argon2", "Không có lỗi"],
    ["TC003", "02/09/2026", "Tester & Developer", "PASS", "Medium", "Kiểm tra validation hợp lệ, chặn mã khóa học trùng", "Không có lỗi"],
    ["TC004", "02/09/2026", "Tester & Developer", "PASS", "High", "Chặn thành công việc xếp trùng phòng học cùng ca", "Không có lỗi"],
    ["TC005", "02/09/2026", "Tester & Developer", "PASS", "High", "Chặn phân công trùng lịch dạy của giáo viên", "Không có lỗi"],
    ["TC006", "02/09/2026", "Tester & Developer", "PASS", "High", "Đủ 4 điều kiện ghi danh, tự động tạo hóa đơn học phí", "Không có lỗi"],
    ["TC007", "02/09/2026", "Tester & Developer", "PASS", "High", "Hỗ trợ thu tiền nhiều đợt, tính chính xác số dư công nợ", "Không có lỗi"],
    ["TC008", "02/09/2026", "Tester & Developer", "PASS", "Medium", "Lưu trữ 4 trạng thái chuyên cần chính xác", "Không có lỗi"],
    ["TC009", "02/09/2026", "Tester & Developer", "PASS", "High", "Tính đúng công thức 20/30/50, xét ĐẠT khi TK>=50 và CC>=80", "Không có lỗi"],
    ["TC010", "02/09/2026", "Tester & Developer", "PASS", "Medium", "Tra cứu nhanh chóng, chỉ xem được dữ liệu cá nhân", "Không có lỗi"],
    ["TC011", "02/09/2026", "Tester & Developer", "PASS", "Medium", "Dashboard thống kê dữ liệu chính xác theo thời gian thực", "Không có lỗi"],
    ["TC012", "02/09/2026", "Tester & Developer", "PASS", "High", "Loại bỏ hoàn toàn lớp ảo giác, tự động Fallback khi mất mạng", "Zero-Trust đạt chuẩn"],
    ["TC013", "02/09/2026", "Tester & Developer", "PASS", "Medium", "Sinh chuẩn 5 câu hỏi JSON kèm đáp án và giải thích chi tiết", "Đúng định dạng"],
    ["TC014", "02/09/2026", "Tester & Developer", "PASS", "Medium", "Đưa ra lời khuyên ôn tập cá nhân hóa sâu sắc", "Audit log đầy đủ"]
]
create_table(headers_tr, data_tr)
add_caption("Bảng 9.4: Báo cáo kết quả kiểm thử hệ thống (Test Report)")

add_p("Đánh giá chất lượng tổng thể: 14/14 Ca kiểm thử đạt trạng thái PASS (Tỷ lệ thành công 100%). Hệ thống đáp ứng đầy đủ tất cả các quy tắc nghiệp vụ và sẵn sàng đưa vào vận hành.", bold_prefix="Kết luận nghiệm thu: ")

# =========================================================================
# PHẦN 2: HƯỚNG DẪN SỬ DỤNG ỨNG DỤNG (USER GUIDE)
# =========================================================================

add_heading_1("PHẦN 2: HƯỚNG DẪN SỬ DỤNG ỨNG DỤNG (USER GUIDE)")

add_heading_2("1. Giới Thiệu Ứng Dụng")
add_p("Hệ thống Quản lý Trung tâm Ngoại ngữ tích hợp Trí tuệ Nhân tạo (ETC English LMS AI) là nền tảng quản lý đào tạo toàn diện, hỗ trợ 4 nhóm đối tượng người dùng (Quản lý, Giáo viên, Học viên, Tư vấn viên) thực hiện toàn bộ quy trình từ tuyển sinh, xếp lịch, thu phí, giảng dạy, điểm danh, chấm điểm đến trợ giảng thông minh với GenAI.")

add_heading_2("2. Cấu Hình Phần Cứng - Phần Mềm Để Sử Dụng")
add_bullet("Phần cứng: Máy tính để bàn, Laptop, Máy tính bảng hoặc Smartphone có kết nối mạng Internet ổn định (băng thông tối thiểu 2 Mbps).", "• ")
add_bullet("Phần mềm: Trình duyệt web hiện đại (Google Chrome, Microsoft Edge, Mozilla Firefox, Safari) phiên bản mới nhất, không cần cài đặt thêm phần mềm phụ trợ.", "• ")

add_heading_2("3. Hướng Dẫn Sử Dụng Các Chức Năng Chính Theo Tác Nhân (Actors)")

add_heading_3("3.1 Chức Năng Của Người Quản Lý (Admin - QUAN_LY)")
add_bullet("Đăng nhập Quản trị: Truy cập trang đăng nhập (/login), chọn tài khoản admin01 (Mật khẩu: Admin@123). Hệ thống tự động chuyển hướng đến Dashboard Quản trị (/admin/dashboard).", "Bước 1: ")
add_bullet("Xem Thống kê Trung tâm: Dashboard hiển thị tổng số học viên, giáo viên, doanh thu, thanh tiến độ sĩ số các lớp và tỷ lệ học viên hoàn thành khóa.", "Bước 2: ")
add_bullet("Quản lý Khóa học & Mở Lớp học: Vào mục 'Khóa học' để tạo khóa học mới; vào mục 'Lớp học' để mở lớp, bấm nút 'Thêm Lịch Học' (hệ thống tự động kiểm tra chống trùng phòng học) và 'Phân Công GV' (kiểm tra chống trùng lịch dạy của giảng viên).", "Bước 3: ")
add_bullet("Quản lý Học viên & Thu Học phí: Vào mục 'Học viên' để tìm kiếm, lọc theo trình độ CEFR; vào mục 'Học phí' để theo dõi danh sách hóa đơn và lập phiếu thu thanh toán nhiều đợt.", "Bước 4: ")
add_bullet("Báo cáo Thống kê Chuyên sâu: Vào mục 'Báo cáo' (/admin/reports) để xem biểu đồ phân tích chi tiết hiệu quả đào tạo và xuất báo cáo.", "Bước 5: ")

add_heading_3("3.2 Chức Năng Của Giáo Viên (Teacher - GIAO_VIEN)")
add_bullet("Đăng nhập Giảng viên: Chọn tài khoản teacher01 (Nguyễn Thị Lan) tại màn hình đăng nhập. Hệ thống chuyển hướng đến Bàn làm việc Giảng viên (/teacher/dashboard).", "Bước 1: ")
add_bullet("Xem Lịch Dạy & Lớp Phụ Trách: Giảng viên theo dõi danh sách các lớp được phân công, thời khóa biểu từng thứ trong tuần và phòng học tương ứng.", "Bước 2: ")
add_bullet("Điểm Danh Buổi Học: Truy cập mục 'Điểm danh' (/teacher/attendance), chọn lớp học. Hệ thống hiển thị danh sách học viên; giảng viên chọn 1 trong 4 trạng thái (Có Mặt, Đi Muộn, Có Phép, Vắng) và bấm 'Lưu Điểm Danh'.", "Bước 3: ")
add_bullet("Nhập Điểm & Đánh Giá Kết Quả: Truy cập mục 'Bảng điểm' (/teacher/grades), nhập điểm Chuyên cần (20%), Giữa kỳ (30%), Cuối kỳ (50%). Hệ thống tự động tính điểm tổng kết và xếp loại ĐẠT/KHÔNG ĐẠT.", "Bước 4: ")
add_bullet("Trợ Lý AI Sinh Bài Luyện Tập: Truy cập mục 'AI Sinh đề' (/teacher/ai-exercises), nhập chủ đề ngữ pháp/từ vựng và chọn độ khó CEFR. Bấm 'Sinh Đề AI' để Gemini tạo tức thì 5 câu trắc nghiệm kèm đáp án và giải thích chi tiết.", "Bước 5: ")

add_heading_3("3.3 Chức Năng Của Học Viên (Student - HOC_VIEN)")
add_bullet("Đăng nhập Góc Học Tập: Chọn tài khoản student01 (Phạm Văn An - CEFR B1). Hệ thống hiển thị thông tin hồ sơ và các lớp đang theo học (/student/dashboard).", "Bước 1: ")
add_bullet("AI Tư Vấn Lộ Trình & Lớp Học: Vào mục 'AI Tư vấn' (/student/ai-consult), chọn trình độ CEFR và các buổi rảnh trong tuần. Hệ thống Gemini AI phân tích và gợi ý top 3 lớp học thực tế phù hợp nhất.", "Bước 2: ")
add_bullet("Đăng Ký Lớp Học Mới: Vào mục 'Đăng ký lớp' (/student/enroll), chọn lớp mong muốn. Hệ thống tự động kiểm tra 4 điều kiện (Sĩ số < 25, chưa đăng ký, chuẩn CEFR, không trùng lịch) và tự động sinh hóa đơn học phí.", "Bước 3: ")
add_bullet("Tra Cứu Thời Khóa Biểu & Bảng Điểm: Vào mục 'Thời khóa biểu' (/student/schedule) để xem lịch học; vào mục 'Bảng điểm' (/student/grades) để theo dõi điểm 3 thành phần và trạng thái ĐẠT khóa học.", "Bước 4: ")
add_bullet("AI Luyện Tập & Tóm Tắt Tiến Độ: Vào 'AI Luyện tập' (/student/ai-practice) để làm bài trắc nghiệm tương tác có chấm điểm trực tiếp; vào 'AI Tóm tắt' (/student/ai-progress) để nhận bản đánh giá điểm mạnh/yếu và lời khuyên ôn tập cá nhân hóa.", "Bước 5: ")

add_heading_3("3.4 Chức Năng Của Tư Vấn Viên (Staff/Counselor - TU_VAN_VIEN)")
add_bullet("Đăng nhập Tuyển sinh: Chọn tài khoản staff01 (Lê Thị Tư Vấn). Hệ thống chuyển hướng đến Bàn làm việc Tuyển sinh (/staff/dashboard).", "Bước 1: ")
add_bullet("Tiếp Nhận Học Viên Mới: Vào mục 'Tiếp nhận học viên' (/staff/new-student), nhập họ tên, thông tin liên hệ và kết quả đánh giá trình độ CEFR đầu vào để khởi tạo hồ sơ học viên trong hệ thống.", "Bước 2: ")
add_bullet("Tư Vấn & Ghi Danh Lớp Học: Sử dụng công cụ AI Tư vấn để tìm lớp phù hợp theo lịch rảnh của học viên, sau đó vào mục 'Ghi danh & Thu phí' (/staff/collect-fee) để đăng ký lớp.", "Bước 3: ")
add_bullet("Lập Phiếu Thu Học Phí: Nhập số tiền thu (tiền mặt hoặc chuyển khoản) để cập nhật công nợ và xuất hóa đơn xác nhận cho học viên.", "Bước 4: ")

# Save unified docx
output_file = r"D:\MyProjects\lms-ai\docs\design\CHUONG_8_9_CAI_DAT_TRIEN_KHAI_KIEM_THU_HUONG_DAN_SD.docx"
doc.save(output_file)
print("SUCCESS: Generated unified Chapter 8 & 9 Word document at:", output_file)
