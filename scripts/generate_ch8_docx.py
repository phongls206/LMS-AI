import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

# Set standard margins (2.54 cm = 1 inch)
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Helper styles
def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def add_heading_1(text):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Dark Blue
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(12)
    return h

def add_heading_2(text):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_heading_3(text):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)
    return h

def add_p(text, bold_prefix="", italic_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Times New Roman'
        r_bold.font.size = Pt(13)
        r_bold.font.bold = True

    if italic_prefix:
        r_it = p.add_run(italic_prefix)
        r_it.font.name = 'Times New Roman'
        r_it.font.size = Pt(13)
        r_it.font.italic = True

    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    return p

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(4)

    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Times New Roman'
        r_bold.font.size = Pt(13)
        r_bold.font.bold = True

    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    return p

def add_code_block(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(30, 41, 59)
    return p

# ----------------- CONTENT -----------------

add_heading_1("CHƯƠNG 8: CÀI ĐẶT VÀ TRIỂN KHAI HỆ THỐNG")

add_heading_2("8.1 Môi trường và Yêu cầu Kỹ thuật")
add_heading_3("8.1.1 Yêu cầu Cấu hình Phần cứng")
add_p("Hệ thống Quản lý trung tâm ngoại ngữ tích hợp AI (ETC English) được xây dựng theo kiến trúc phân tán (Multi-tier Cloud-ready Architecture), cho phép tối ưu hóa tài nguyên phần cứng máy chủ và máy khách, giảm thiểu chi phí đầu tư hạ tầng ban đầu:")

# Table 8.1
table1 = doc.add_table(rows=4, cols=3)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
table1.style = 'Table Grid'

headers1 = ["Thành phần", "Cấu hình tối thiểu", "Cấu hình khuyến nghị"]
for idx, text in enumerate(headers1):
    cell = table1.cell(0, idx)
    set_cell_background(cell, "1E3A8A")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

data1 = [
    ["Máy chủ Phát triển (Dev Server)", "CPU 2 Cores, RAM 4GB, Ổ cứng trống 5GB", "CPU 4 Cores, RAM 8GB–16GB, Ổ cứng SSD 20GB"],
    ["Máy khách Người dùng (Client)", "Thiết bị có trình duyệt web (PC, Laptop, Smartphone)", "PC/Laptop màn hình Full HD, kết nối Internet ổn định"],
    ["Băng thông Mạng", "Tối thiểu 5 Mbps (truy vấn và gọi GenAI API)", "Khuyến nghị 20 Mbps trở lên để tải trang mượt mà"]
]

for row_idx, row_data in enumerate(data1, start=1):
    for col_idx, val in enumerate(row_data):
        cell = table1.cell(row_idx, col_idx)
        if row_idx % 2 == 1:
            set_cell_background(cell, "F8FAFC")
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

caption1 = doc.add_paragraph()
caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_cap1 = caption1.add_run("Bảng 8.1: Yêu cầu cấu hình phần cứng tối thiểu và khuyến nghị")
r_cap1.font.name = 'Times New Roman'
r_cap1.font.italic = True
r_cap1.font.size = Pt(11)
caption1.paragraph_format.space_after = Pt(10)

add_heading_3("8.1.2 Yêu cầu Môi trường Phần mềm & Công nghệ")
add_p("Hệ thống được phát triển hoàn toàn trên nền tảng TypeScript xuyên suốt từ Frontend đến Backend:")

# Table 8.2
table2 = doc.add_table(rows=10, cols=4)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.style = 'Table Grid'

headers2 = ["Thành phần", "Công nghệ / Framework", "Phiên bản", "Mục đích sử dụng"]
for idx, text in enumerate(headers2):
    cell = table2.cell(0, idx)
    set_cell_background(cell, "1E3A8A")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

data2 = [
    ["Runtime Môi trường", "Node.js (LTS)", "v20.x / v22.x", "Môi trường thực thi JavaScript/TypeScript phía máy chủ"],
    ["Quản lý Gói", "npm", "10.x+", "Quản lý và cài đặt các thư viện phụ thuộc"],
    ["Backend Framework", "NestJS", "v11.x / v12.x", "Xây dựng RESTful API kiến trúc phân tầng Modular"],
    ["ORM / Data Access Layer", "Prisma ORM", "v6.4.1 (Stable)", "Quản lý Schema, Migration và Type-safe Database Client"],
    ["Frontend Framework", "Next.js (App Router)", "v14.x / v16.x", "Xây dựng giao diện SSR/CSR với TypeScript"],
    ["CSS & Giao diện", "TailwindCSS + Lucide Icons", "v3.4.x / v4.x", "Thiết kế giao diện hiện đại, Responsive đa thiết bị"],
    ["Cơ sở Dữ liệu", "PostgreSQL (Neon.tech)", "v15+ / v16+", "CSDL quan hệ Serverless Cloud Database"],
    ["Trí tuệ Nhân tạo", "Google Gemini SDK", "@google/genai", "Tích hợp mô hình ngôn ngữ lớn (Gemini 2.5 Flash/Pro)"],
    ["Bảo mật & Mã hóa", "Argon2 + JWT", "argon2, @nestjs/jwt", "Băm mật khẩu an toàn và xác thực phân quyền Stateless"]
]

for row_idx, row_data in enumerate(data2, start=1):
    for col_idx, val in enumerate(row_data):
        cell = table2.cell(row_idx, col_idx)
        if row_idx % 2 == 1:
            set_cell_background(cell, "F8FAFC")
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

caption2 = doc.add_paragraph()
caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_cap2 = caption2.add_run("Bảng 8.2: Danh mục công nghệ và phiên bản phần mềm cốt lõi")
r_cap2.font.name = 'Times New Roman'
r_cap2.font.italic = True
r_cap2.font.size = Pt(11)
caption2.paragraph_format.space_after = Pt(10)

add_heading_3("8.1.3 Cấu trúc Thư mục Mã Nguồn Chuẩn")
add_p("Dự án được tổ chức theo mô hình Monorepo rõ ràng giữa Backend và Frontend:")
add_code_block("""D:\\MyProjects\\lms-ai\\
├── backend/                        # Máy chủ Backend (NestJS 12)
│   ├── prisma/
│   │   ├── schema.prisma           # 14 Models thực thể chuẩn 3NF
│   │   └── seed.ts                 # Script nạp dữ liệu mẫu ban đầu
│   ├── src/
│   │   ├── common/                 # Guards, Decorators, Strategies (RBAC)
│   │   ├── modules/
│   │   │   ├── auth/               # Module xác thực JWT & Argon2
│   │   │   ├── users/              # Module học viên & giáo viên
│   │   │   ├── courses/            # Module quản lý khóa học
│   │   │   ├── classes/            # Module lớp học & lịch học (chống trùng phòng/GV)
│   │   │   ├── enrollments/        # Module đăng ký lớp & hóa đơn học phí
│   │   │   ├── attendances/        # Module buổi học & điểm danh 4 trạng thái
│   │   │   ├── grades/             # Module bảng điểm (công thức 20/30/50)
│   │   │   ├── statistics/         # Module báo cáo doanh thu & tỷ lệ hoàn thành
│   │   │   └── ai/                 # Module tích hợp Gemini AI & Lọc ảo giác
│   │   ├── prisma/                 # Prisma Module & Prisma Service
│   │   ├── app.module.ts           # Root Module của NestJS
│   │   └── main.ts                 # Điểm khởi chạy hệ thống, Swagger & Validation
│   └── .env                        # File cấu hình biến môi trường Backend
├── frontend/                       # Ứng dụng Giao diện (Next.js 14 App Router)
│   ├── src/
│   │   ├── app/                    # 23 Màn hình phân luồng theo 4 vai trò (RBAC)
│   │   ├── components/             # AppLayout, Sidebar, Header phân quyền
│   │   ├── services/               # Axios API Client với JWT Interceptor
│   │   └── types/                  # TypeScript Data Models & DTO Interfaces
└── docs/                           # Tài liệu thiết kế hệ thống & Đồ án""")

add_heading_2("8.2 Cấu Hình Môi Trường & Quản Trị Bí Mật (Environment & Secrets)")
add_heading_3("8.2.1 Cấu hình Biến Môi Trường Backend (backend/.env)")
add_p("Toàn bộ thông tin cấu hình nhạy cảm được quản lý qua biến môi trường độc lập, tuyệt đối không hardcode vào mã nguồn:")
add_code_block("""# Cổng dịch vụ Backend
PORT=8000

# Chuỗi kết nối Neon Serverless PostgreSQL (Pooler Mode hỗ trợ SSL)
DATABASE_URL="postgresql://neondb_owner:npg_3lFnjQKIo5rM@ep-dark-resonance-axjf1mzr-pooler.c-4.us-east-2.aws.neon.tech/neondb?sslmode=require&channel_binding=require"

# Khóa bí mật ký phát mã JWT (Stateless Token)
JWT_SECRET="etc_english_center_jwt_secret_key_2026_super_secure"
JWT_EXPIRES_IN="24h"

# Khóa API Google Gemini AI (Tích hợp AI)
GEMINI_API_KEY="AIzaSy...your_gemini_api_key_here..."

# Cấu hình CORS cho phép Frontend truy cập
FRONTEND_URL="http://localhost:3000" """)

add_heading_3("8.2.2 Cấu hình Chuỗi Kết Nối Cơ Sở Dữ Liệu Neon Cloud")
add_p("Hệ thống sử dụng đường dẫn kết nối dạng Connection Pooler (PgBouncer) của Neon.tech để tối ưu số lượng kết nối đồng thời từ NestJS, đồng thời bật chế độ mã hóa đường truyền bắt buộc (sslmode=require).")

add_heading_2("8.3 Quy Trình Cài Đặt Và Khởi Tạo Cơ Sở Dữ Liệu")
add_p("Quy trình thiết lập hệ thống từ mã nguồn được thực hiện theo 3 bước tuần tự:")
add_bullet("Cài đặt thư viện phụ thuộc (Dependencies): Chạy npm install tại thư mục backend và frontend.", "Bước 1: ")
add_bullet("Đồng bộ CSDL và Sinh Prisma Client: Chạy npx prisma db push và npx prisma generate để tạo 14 bảng quan hệ 3NF trên Neon PostgreSQL.", "Bước 2: ")
add_bullet("Nạp Dữ liệu Mẫu Ban Đầu (Database Seeding): Chạy npm run db:seed để nạp 4 vai trò người dùng, 2 khóa học, 2 lớp học và lịch học.", "Bước 3: ")

# Table 8.3
table3 = doc.add_table(rows=7, cols=5)
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
table3.style = 'Table Grid'

headers3 = ["Tên đăng nhập", "Mật khẩu", "Vai trò (RBAC)", "Họ và Tên", "Mô tả nghiệp vụ"]
for idx, text in enumerate(headers3):
    cell = table3.cell(0, idx)
    set_cell_background(cell, "1E3A8A")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

data3 = [
    ["admin01", "Admin@123", "QUAN_LY", "Nguyễn Quản Lý", "Toàn quyền quản trị trung tâm, tài chính, phân công"],
    ["teacher01", "Admin@123", "GIAO_VIEN", "Nguyễn Thị Lan", "Giảng viên IELTS, TOEIC (Phụ trách lớp LOP01)"],
    ["teacher02", "Admin@123", "GIAO_VIEN", "Trần Văn Minh", "Giảng viên Giao tiếp (Phụ trách lớp LOP02)"],
    ["staff01", "Admin@123", "TU_VAN_VIEN", "Lê Thị Tư Vấn", "Tư vấn viên tiếp nhận học viên & thu phí tại quầy"],
    ["student01", "Admin@123", "HOC_VIEN", "Phạm Văn An", "Học viên trình độ CEFR B1 (Lớp LOP01)"],
    ["student02", "Admin@123", "HOC_VIEN", "Hoàng Thị Bình", "Học viên trình độ CEFR A2 (Lớp LOP02)"]
]

for row_idx, row_data in enumerate(data3, start=1):
    for col_idx, val in enumerate(row_data):
        cell = table3.cell(row_idx, col_idx)
        if row_idx % 2 == 1:
            set_cell_background(cell, "F8FAFC")
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

caption3 = doc.add_paragraph()
caption3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_cap3 = caption3.add_run("Bảng 8.3: Danh mục tài khoản người dùng mẫu sau khi nạp Seeding")
r_cap3.font.name = 'Times New Roman'
r_cap3.font.italic = True
r_cap3.font.size = Pt(11)
caption3.paragraph_format.space_after = Pt(10)

add_heading_2("8.4 Quy Trình Khởi Chạy Hệ Thống Trên Môi Trường Cục Bộ")
add_bullet("Khởi chạy Máy chủ Backend NestJS: cd backend && npm run start:dev (Backend chạy tại port 8000, Swagger UI tại http://localhost:8000/api/docs).", "1. ")
add_bullet("Khởi chạy Ứng dụng Giao diện Frontend Next.js: cd frontend && npm run dev (Giao diện chạy tại http://localhost:3000/login).", "2. ")
add_bullet("Quản lý Dữ liệu Trực quan Qua Prisma Studio: cd backend && npx prisma studio (Giao diện GUI quản lý 14 bảng CSDL tại http://localhost:5555).", "3. ")

add_heading_2("8.5 Phương Án Triển Khai Lên Môi Trường Đám Mây (Cloud Deployment)")
add_p("Hệ thống được thiết kế sẵn sàng cho việc triển khai phân tán không máy chủ (Serverless Cloud Architecture):")
add_bullet("Frontend Tier: Triển khai trên nền tảng Vercel (Edge Network / Serverless Functions) với tên miền chính.", "• ")
add_bullet("Backend Tier: Triển khai trên Render hoặc Railway với môi trường Docker Node.js Container.", "• ")
add_bullet("Database Tier: Lưu trữ trực tiếp trên Neon Serverless PostgreSQL với tính năng tự động sao lưu và mở rộng quy mô (Auto-scaling).", "• ")
add_bullet("Chính sách Bảo mật: Bật tường lửa CORS giới hạn Domain truy cập, áp dụng HTTPS toàn bộ luồng dữ liệu và kích hoạt Rate Limiting ngăn chặn tấn công DDoS.", "• ")

# Save document
output_path = r"D:\MyProjects\lms-ai\docs\design\CHUONG_8_CAI_DAT_VA_TRIEN_KHAI_HE_THONG.docx"
doc.save(output_path)
print("Saved DOCX successfully at:", output_path)
